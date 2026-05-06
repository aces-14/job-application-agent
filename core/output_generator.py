import io
import os
import json
from datetime import datetime
from typing import List, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.resume_tailor import TailoredResume, ChangeEntry
from core.validator import IntegrityReport
from core.config import OUTPUT_DIR
from core.logger import setup_logger

logger = setup_logger(__name__)


# ---------------------------------------------------------------------------
# Template-based helpers (preserve original formatting)
# ---------------------------------------------------------------------------

def _all_paragraphs(doc: Document):
    """Yield every paragraph in the document, including those inside tables."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _set_para_text(para, new_text: str) -> None:
    """
    Replace a paragraph's text in-place, keeping the first run's character
    formatting (font, size, bold) for the entire replacement.
    Subsequent runs are cleared so text isn't doubled.
    """
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _apply_changes_to_template(
    tailored: TailoredResume,
    original_bytes: bytes,
    path: str,
) -> bool:
    """
    Load the original .docx, apply each changelog change in-place, and save.
    Returns True on success, False if the template approach should be skipped.
    """
    try:
        doc = Document(io.BytesIO(original_bytes))
        paras = list(_all_paragraphs(doc))
        applied = 0

        for change in tailored.changelog:
            if not change.original or not change.revised:
                continue
            needle = change.original.strip()
            for para in paras:
                if needle in para.text:
                    new_text = para.text.replace(needle, change.revised.strip())
                    _set_para_text(para, new_text)
                    applied += 1
                    break
                # Partial match — try first 50 chars (LLM sometimes truncates original)
                if len(needle) > 50 and needle[:50] in para.text:
                    _set_para_text(para, change.revised.strip())
                    applied += 1
                    break

        doc.save(path)
        logger.info(
            f"Template resume saved: {path} "
            f"({applied}/{len(tailored.changelog)} changes applied)"
        )
        return True

    except Exception as e:
        logger.warning(f"Template approach failed ({e}), falling back to rebuild")
        return False


# ---------------------------------------------------------------------------
# Fallback: rebuilt from scratch (used when no original bytes are available)
# ---------------------------------------------------------------------------

def _section_heading(doc: Document, text: str) -> None:
    """Add a bold, slightly-larger section heading with a bottom border line."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)  # dark navy

    # Thin horizontal rule below the heading
    hr = doc.add_paragraph("─" * 72)
    hr.runs[0].font.size = Pt(7)
    hr.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    hr.paragraph_format.space_after = Pt(4)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)


def _normal(doc: Document, text: str, size: int = 10, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Public generators
# ---------------------------------------------------------------------------

def generate_resume_docx(
    tailored: TailoredResume,
    session_id: str,
    original_docx_bytes: Optional[bytes] = None,
) -> str:
    """
    Produce the tailored resume as a Word document.

    If original_docx_bytes is supplied, changes are applied in-place onto the
    original .docx so the candidate's formatting is fully preserved.
    Falls back to a clean rebuild if the template approach fails.

    Returns the file path.
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = f"{OUTPUT_DIR}/resume_{session_id}.docx"

    # Prefer template approach — preserves candidate's original formatting
    if original_docx_bytes:
        if _apply_changes_to_template(tailored, original_docx_bytes, path):
            return path

    # Fallback: rebuild from scratch
    doc = Document()

    # ── Narrow margins for a clean look ──────────────────────────────────
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # ── Name ─────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(tailored.full_name)
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(4)

    # ── Contact line ──────────────────────────────────────────────────────
    contact_parts = [
        tailored.email,
        tailored.phone,
        tailored.location,
        tailored.linkedin,
        tailored.github,
    ]
    contact_str = "  |  ".join(p for p in contact_parts if p)
    if contact_str:
        c = doc.add_paragraph(contact_str)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].font.size = Pt(9)
        c.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        c.paragraph_format.space_after = Pt(8)

    # ── Summary ───────────────────────────────────────────────────────────
    if tailored.summary:
        _section_heading(doc, "Professional Summary")
        _normal(doc, tailored.summary)
        doc.add_paragraph("")

    # ── Skills ────────────────────────────────────────────────────────────
    if tailored.skills:
        _section_heading(doc, "Skills")
        # Four skills per row, separated by bullet
        rows = [tailored.skills[i:i + 4] for i in range(0, len(tailored.skills), 4)]
        for row in rows:
            _normal(doc, "  •  ".join(row))
        doc.add_paragraph("")

    # ── Work Experience ───────────────────────────────────────────────────
    if tailored.work_experience:
        _section_heading(doc, "Work Experience")
        for exp in tailored.work_experience:
            title = exp.get("title", "")
            company = exp.get("company", "")
            duration = exp.get("duration", "")

            role_para = doc.add_paragraph()
            role_run = role_para.add_run(f"{title}  —  {company}")
            role_run.bold = True
            role_run.font.size = Pt(11)
            role_para.paragraph_format.space_after = Pt(1)

            if duration:
                _normal(doc, duration, size=9, italic=True)

            for resp in exp.get("responsibilities", []):
                _bullet(doc, resp)
            for ach in exp.get("achievements", []):
                _bullet(doc, f"✓  {ach}")

            doc.add_paragraph("")

    # ── Education ─────────────────────────────────────────────────────────
    if tailored.education:
        _section_heading(doc, "Education")
        for edu in tailored.education:
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            institution = edu.get("institution", "")
            year = edu.get("year", "")
            gpa = edu.get("gpa")

            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(f"{degree} in {field}  —  {institution}")
            edu_run.bold = True
            edu_run.font.size = Pt(11)

            details = year
            if gpa:
                details += f"  |  GPA: {gpa}"
            if details:
                _normal(doc, details, size=9, italic=True)

            doc.add_paragraph("")

    # ── Certifications ────────────────────────────────────────────────────
    if tailored.certifications:
        _section_heading(doc, "Certifications")
        for cert in tailored.certifications:
            _bullet(doc, cert)
        doc.add_paragraph("")

    # ── Projects ──────────────────────────────────────────────────────────
    if tailored.projects:
        _section_heading(doc, "Projects")
        for proj in tailored.projects:
            _bullet(doc, proj)
        doc.add_paragraph("")

    # ── Languages ─────────────────────────────────────────────────────────
    if tailored.languages:
        _section_heading(doc, "Languages")
        _normal(doc, "  •  ".join(tailored.languages))

    doc.save(path)
    logger.info(f"Resume docx saved: {path}")
    return path


def generate_text_outputs(
    cover_letter: str,
    outreach_message: str,
    session_id: str,
) -> tuple[str, str]:
    """Save cover letter and outreach as plain text. Returns (cl_path, out_path)."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    cl_path = f"{OUTPUT_DIR}/cover_letter_{session_id}.txt"
    out_path = f"{OUTPUT_DIR}/outreach_{session_id}.txt"

    with open(cl_path, "w", encoding="utf-8") as f:
        f.write(cover_letter)

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(outreach_message)

    logger.info(f"Text outputs saved: {cl_path}, {out_path}")
    return cl_path, out_path


def generate_match_report(
    match_matrix: dict,
    integrity_report: IntegrityReport,
    changelog: List[ChangeEntry],
    session_id: str,
) -> str:
    """Save the full match + integrity + changelog report as JSON. Returns file path."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = f"{OUTPUT_DIR}/match_report_{session_id}.json"

    report = {
        "session_id": session_id,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "match_matrix": match_matrix,
        "integrity": {
            "passed": integrity_report.passed,
            "violations": integrity_report.violations,
            "not_added": integrity_report.not_added,
        },
        "changelog": [
            {
                "section": c.section,
                "original": c.original,
                "revised": c.revised,
                "reason": c.reason,
                "change_type": c.change_type,
                "confidence": c.confidence,
                "confidence_reason": c.confidence_reason,
                "ats_keywords_added": c.ats_keywords_added,
            }
            for c in changelog
        ],
        "confidence_summary": {
            "high": sum(1 for c in changelog if c.confidence == "high"),
            "medium": sum(1 for c in changelog if c.confidence == "medium"),
            "low": sum(1 for c in changelog if c.confidence == "low"),
        },
    }

    with open(path, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2, ensure_ascii=False)

    logger.info(f"Match report saved: {path}")
    return path
