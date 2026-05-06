import re
import json
import io
from typing import Optional, List
from docx import Document
from pydantic import BaseModel, Field
from core.logger import setup_logger
from groq import Groq
from core.config import GROQ_API_KEY, GROQ_MODEL, GROQ_MAX_TOKENS

logger = setup_logger(__name__)

class WorkExperience(BaseModel):
    company: str
    title: str
    duration: str
    responsibilities: List[str]
    achievements: List[str]

class Education(BaseModel):
    institution: str
    degree: Optional[str] = None
    field: Optional[str] = None
    year: Optional[str] = None
    gpa: Optional[str] = None

class ResumeData(BaseModel):
    full_name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    summary: Optional[str] = None
    skills: List[str] = Field(default_factory=list)
    work_experience: List[WorkExperience] = Field(default_factory=list)
    education: List[Education] = Field(default_factory=list)
    certifications: List[str] = Field(default_factory=list)
    projects: List[str] = Field(default_factory=list)
    languages: List[str] = Field(default_factory=list)
    raw_text: str = ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables are common in resumes for skills / layout columns
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))

        text = "\n".join(parts)
        logger.info(f"Extracted {len(text)} characters from DOCX")
        return text
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise ValueError(f"Could not extract text from DOCX: {str(e)}")

def parse_resume_with_llm(raw_text: str) -> ResumeData:
    client = Groq(api_key=GROQ_API_KEY)

    prompt = f"""Extract all information from this resume and return ONLY valid JSON.
Do NOT add any information not explicitly present in the resume.
Do NOT infer or assume skills, experience, or achievements.
Extract EXACTLY what is written — nothing more, nothing less.

Return this exact JSON structure:
{{
    "full_name": "string",
    "email": "string or null",
    "phone": "string or null",
    "location": "string or null",
    "linkedin": "string or null",
    "github": "string or null",
    "summary": "string or null",
    "skills": ["list of skills EXACTLY as written"],
    "work_experience": [
        {{
            "company": "string",
            "title": "string",
            "duration": "string",
            "responsibilities": ["list exactly as written"],
            "achievements": ["list exactly as written"]
        }}
    ],
    "education": [
        {{
            "institution": "string",
            "degree": "string",
            "field": "string",
            "year": "string",
            "gpa": "string or null"
        }}
    ],
    "certifications": ["list exactly as written"],
    "projects": ["list exactly as written"],
    "languages": ["list exactly as written"]
}}

RESUME TEXT:
{raw_text}

RULES:
- Return ONLY the JSON object, no explanation
- If a field is not present, use null or empty list
- Do NOT add skills that are implied but not stated
- Do NOT expand abbreviations unless the full form is in the resume
- Copy text EXACTLY as written, preserving the candidate's own words"""

    try:
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=GROQ_MAX_TOKENS,
            temperature=0.0
        )

        content = response.choices[0].message.content.strip()

        # Clean up markdown code blocks if present
        if content.startswith("```"):
            content = re.sub(r'^```(?:json)?\n?', '', content)
            content = re.sub(r'\n?```$', '', content)

        parsed = json.loads(content)
        resume_data = ResumeData(**parsed, raw_text=raw_text)
        logger.info(
            f"Resume parsed successfully: {resume_data.full_name}, "
            f"{len(resume_data.skills)} skills, "
            f"{len(resume_data.work_experience)} positions"
        )
        return resume_data

    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing failed: {str(e)}")
        raise ValueError(f"LLM returned invalid JSON: {str(e)}")
    except Exception as e:
        logger.error(f"Resume parsing failed: {str(e)}")
        raise

def parse_resume(file_bytes: bytes) -> ResumeData:
    logger.info("Starting resume parsing pipeline")

    raw_text = extract_text_from_docx(file_bytes)

    if len(raw_text) < 100:
        raise ValueError(
            "Resume content is too short. "
            "Please ensure the .docx file contains your full resume."
        )

    resume_data = parse_resume_with_llm(raw_text)

    if not resume_data.full_name or resume_data.full_name == "string":
        raise ValueError(
            "Could not extract candidate name from resume. "
            "Please check the .docx file."
        )

    return resume_data